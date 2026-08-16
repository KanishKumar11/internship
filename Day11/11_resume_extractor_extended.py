"""
11_resume_extractor_extended.py - The three extensions

TEACHES : The extension ideas from the deck, in one file - pull a skills
          list out of the text, accept .docx as well as .pdf, and wrap the
          whole thing in a Streamlit upload form. The shape the group
          project will take in Phase 4.
SLIDE   : Day 11, Slide 19 - Exercise, Extend (deck page 19/20)
RUN     : streamlit run 11_resume_extractor_extended.py
          python 11_resume_extractor_extended.py   (falls back to a
          terminal run against sample_resume.pdf, so you can check the
          logic without a browser)

EXPECTED OUTPUT
    In the browser: an upload box that takes a .pdf or .docx, then Name,
    Email, Phone, the detected skills as a list, and the raw extracted
    text in an expander.
    In the terminal (no Streamlit): the same fields printed for
    sample_resume.pdf -
        Name  : Aarav Sharma
        Email : aarav.sharma@example.com
        Phone : +91 98765 43210
        Skills: Python, JavaScript, SQL, SQLite, HTML, CSS, Streamlit,
                pandas, Git
    The skills come out in KNOWN_SKILLS order, not resume order - the
    loop walks the known list and asks "is this one in the text?"

REQUIRES
    sample_resume.pdf - run create_sample_resume.py first.
    pip install streamlit    (only for the browser version)

ONLY OPEN THIS AFTER FILE 10 WORKS. Core extraction first, extras second.
"""

import re
from pathlib import Path

try:
    from PyPDF2 import PdfReader
except ImportError:
    from pypdf import PdfReader

SAMPLE_PDF = Path(__file__).with_name("sample_resume.pdf")

EMAIL_PATTERN = r"[\w.+-]+@[\w-]+\.[\w.]+"
PHONE_PATTERN = r"\+?\d[\d\s-]{8,}\d"

# The skills we know how to spot. A real matcher would load this from a
# database or a job description - which is exactly what Day 13 does.
KNOWN_SKILLS: list[str] = [
    "Python", "Java", "JavaScript", "C++", "SQL", "SQLite",
    "HTML", "CSS", "React", "Streamlit", "pandas", "NumPy", "Git", "Excel",
]


# === EXTENSION 2 === Handle Word docs too ===============================
# One function, two formats. Everything downstream works on a plain
# string, so the rest of the file never learns which kind of file it was.
def extract_text(source: object, filename: str) -> str:
    """Extract text from a .pdf or .docx. `source` is a path or a file object."""
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        reader = PdfReader(source)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        return text

    if suffix == ".docx":
        from docx import Document

        doc = Document(source)
        # Paragraphs first, then table cells - a resume in a table layout
        # would otherwise come back empty, and those do exist.
        text = "\n".join(para.text for para in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)
        return text

    # Fail loudly on anything else rather than returning "" and letting
    # the caller wonder why every pattern found nothing.
    raise ValueError(f"Cannot read {suffix or 'a file with no extension'} - use .pdf or .docx")


def find_email(text: str) -> str:
    """The first email address in the text."""
    matches = re.findall(EMAIL_PATTERN, text)
    return matches[0] if matches else "Not found"


def find_phone(text: str) -> str:
    """The first phone number in the text."""
    matches = re.findall(PHONE_PATTERN, text)
    return matches[0] if matches else "Not found"


def find_name(text: str) -> str:
    """The first non-empty line - on a resume, that is the name."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return lines[0] if lines else "Unknown"


# === EXTENSION 1 === Extract skills =====================================
def find_skills_section(text: str, lines_after: int = 5) -> str:
    """Return the few lines that follow the 'Skills' heading."""
    lines = text.split("\n")
    for index, line in enumerate(lines):
        # .lower() so this works whether the resume says SKILLS, Skills or
        # "Technical Skills". Matching the heading is the fragile part of
        # this whole idea - resumes have no standard section names.
        if "skill" in line.lower():
            return "\n".join(lines[index + 1 : index + 1 + lines_after])
    return ""


def find_skills(text: str) -> list[str]:
    """Find known skill keywords, preferring the Skills section."""
    # Search the Skills section if there is one, the whole document if not.
    # Searching everything would list "Python" for someone who only
    # mentioned it in a hobby line - narrowing to the section first is
    # what makes the result mean something.
    section = find_skills_section(text)
    haystack = section if section.strip() else text

    found: list[str] = []
    for skill in KNOWN_SKILLS:
        # re.escape because "C++" contains two regex quantifiers - without
        # it the pattern would be invalid. \b stops "SQL" matching inside
        # "SQLite", and IGNORECASE catches "python" written in lower case
        # while the canonical spelling is what gets appended.
        if re.search(rf"\b{re.escape(skill)}\b", haystack, re.IGNORECASE):
            found.append(skill)
    return found


def analyse(text: str) -> dict[str, object]:
    """Run every extractor and return the results in one dict."""
    return {
        "name": find_name(text),
        "email": find_email(text),
        "phone": find_phone(text),
        "skills": find_skills(text),
    }


# === EXTENSION 3 === Streamlit UI =======================================
def run_streamlit_app() -> None:
    """The browser version: upload a resume, see the fields."""
    import streamlit as st

    st.title("Resume Contact Extractor")
    st.write("Upload a resume as a **.pdf** or **.docx** and pull the contact details out of it.")

    # st.file_uploader hands back a file-like object, NOT a path - the file
    # never touches the disk. That is why extract_text takes a `source`
    # object plus a separate filename, instead of just a path.
    uploaded = st.file_uploader("Resume", type=["pdf", "docx"], key="resume_upload")

    if uploaded is None:
        st.info("No file yet. Upload sample_resume.pdf from this folder to try it.")
        return

    try:
        text = extract_text(uploaded, uploaded.name)
    except ValueError as error:
        st.error(str(error))
        return

    if not text.strip():
        # The scanned-PDF case from slide 13. Say so plainly rather than
        # showing three empty fields and letting the user think it broke.
        st.warning(
            "No text came out of that file. If it is a scanned PDF, the page "
            "is an image and there is nothing for PyPDF2 to read."
        )
        return

    results = analyse(text)

    col_left, col_right = st.columns(2)
    col_left.metric("Name", results["name"])
    col_right.metric("Phone", results["phone"])
    st.write(f"**Email:** {results['email']}")

    skills = results["skills"]
    if skills:
        st.write(f"**Skills found ({len(skills)}):** {', '.join(skills)}")
    else:
        st.write("**Skills found:** none of the known keywords appeared.")

    with st.expander("Raw extracted text"):
        # Always give the user a way to see the text the patterns ran
        # against. When an extraction looks wrong, this is the first
        # thing to check.
        st.text(text)


def run_in_terminal() -> None:
    """The no-browser version, so the logic can be checked without Streamlit."""
    if not SAMPLE_PDF.exists():
        print(f"{SAMPLE_PDF.name} is missing. Run: python create_sample_resume.py")
        return

    text = extract_text(str(SAMPLE_PDF), SAMPLE_PDF.name)
    results = analyse(text)

    print(f"Name  : {results['name']}")
    print(f"Email : {results['email']}")
    print(f"Phone : {results['phone']}")
    print(f"Skills: {', '.join(results['skills'])}")
    print("\nFor the upload form: streamlit run 11_resume_extractor_extended.py")


try:
    import streamlit  # noqa: F401  (imported to test availability)

    run_streamlit_app()
except ImportError:
    # Streamlit is not installed - fall back to the terminal version so
    # the extraction logic can still be demonstrated and checked.
    run_in_terminal()
