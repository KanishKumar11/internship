"""
create_sample_docx.py - Build the sample_report.docx used by file 07

TEACHES : Nothing - this is a setup script. Run it once before class so
          every laptop has the same Word document to extract from.
SLIDE   : none (it builds the file slide 14 refers to)
RUN     : python create_sample_docx.py

EXPECTED OUTPUT IN THE TERMINAL
        Wrote sample_report.docx
        Check: 2 headings, 5 body paragraphs, 1 table (4 rows x 3 cols)
        Check: email found -> anjali.kaur@college.edu

REQUIRES
    pip install python-docx
    (the import is "docx", not "python_docx" - a classic first-run error)

WHY A TABLE AND HEADINGS
    Slide 14's whole argument is that a .docx keeps structure a PDF loses.
    That comparison only lands if the sample file HAS structure - so this
    document deliberately contains real Heading 1 styles and a real table.
"""

from pathlib import Path

DOCX_FILE = Path(__file__).with_name("sample_report.docx")

# (name, branch, marks) - the table rows, after the header row.
STUDENT_ROWS: list[tuple[str, str, str]] = [
    ("Aarav", "BCA", "85"),
    ("Priya", "BCA", "92"),
    ("Rahul", "BSc-IT", "78"),
]


def build_document(path: Path) -> None:
    """Write a .docx with headings, paragraphs and a table."""
    from docx import Document

    doc = Document()

    # add_heading applies the built-in "Heading 1" style, which is what
    # file 07 looks for with para.style.name == "Heading 1".
    doc.add_heading("Semester III Progress Report", level=1)

    doc.add_paragraph("Hindu College Amritsar - BCA Semester III")
    doc.add_paragraph("Report generated on 2026-08-07 for internal circulation.")
    # An email in the body, so the regex step at the end of file 07 has
    # something to find - the same pattern that worked on the PDF.
    doc.add_paragraph(
        "Queries about this report go to anjali.kaur@college.edu "
        "or call +91 98765 43210."
    )

    doc.add_heading("Class Performance", level=1)
    doc.add_paragraph("Marks for the three students profiled this term:")

    # A real table: 4 rows (1 header + 3 students) x 3 columns.
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"  # visible borders, so it looks like a table
    header = table.rows[0].cells
    header[0].text = "Name"
    header[1].text = "Branch"
    header[2].text = "Marks"

    for name, branch, marks in STUDENT_ROWS:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = branch
        cells[2].text = marks

    doc.add_paragraph("Prepared by the Department of Computer Applications.")

    doc.save(str(path))


def verify(path: Path) -> None:
    """Read the document back and report what file 07 will see in it."""
    from docx import Document

    doc = Document(str(path))

    headings = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
    # A heading is a paragraph too, so the body count has to exclude them.
    # Filter on the STYLE, not on "p not in headings" - python-docx builds
    # a fresh wrapper object every time you touch doc.paragraphs, so two
    # wrappers around the same paragraph are never equal to each other.
    body = [
        p for p in doc.paragraphs if p.text.strip() and p.style.name != "Heading 1"
    ]

    table = doc.tables[0]
    print(
        f"Check: {len(headings)} headings, {len(body)} body paragraphs, "
        f"{len(doc.tables)} table ({len(table.rows)} rows x {len(table.columns)} cols)"
    )

    full_text = "\n".join(p.text for p in doc.paragraphs)
    found = "anjali.kaur@college.edu" in full_text
    print(f"Check: email found -> anjali.kaur@college.edu ... {'OK' if found else 'FAILED'}")


try:
    build_document(DOCX_FILE)
    print(f"Wrote {DOCX_FILE.name}")
    verify(DOCX_FILE)
except ImportError:
    # python-docx has no pure-Python fallback worth writing - a .docx is a
    # zip of XML files, which is a lot of scaffolding for a sample file.
    print("python-docx is not installed, so sample_report.docx was not created.")
    print("Install it with:  pip install python-docx")
    print("File 07 will report the missing file rather than crash.")
