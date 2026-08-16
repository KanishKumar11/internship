"""
07_docx_extraction.py - Pulling text out of a .docx with python-docx

TEACHES : The same five steps as the PDF demo, plus the two things a Word
          document keeps that a PDF throws away - real tables and real
          heading styles.
SLIDE   : Day 11, Slide 14 - Extraction, Word Docs with python-docx
          (deck page 14/20)
RUN     : python 07_docx_extraction.py

EXPECTED OUTPUT IN THE TERMINAL
        --- 1. PARAGRAPHS (7 of them) ---
        Semester III Progress Report
        Hindu College Amritsar - BCA Semester III
        ...
        --- 2. TABLES (1) ---
        Table 1: 4 rows x 3 columns
          ['Name', 'Branch', 'Marks']
          ['Aarav', 'BCA', '85']
          ['Priya', 'BCA', '92']
          ['Rahul', 'BSc-IT', '78']
        --- 3. HEADINGS ---
        HEADING: Semester III Progress Report
        HEADING: Class Performance
        --- 4. REGEX ON THE TEXT ---
        Emails found: ['anjali.kaur@college.edu']

REQUIRES
    pip install python-docx      (the import is "docx", not "python_docx")
    sample_report.docx - run create_sample_docx.py first.
"""

import re
from pathlib import Path

DOCX_FILE = Path(__file__).with_name("sample_report.docx")

EMAIL_PATTERN = r"[\w.+-]+@[\w-]+\.[\w.]+"

# The install name and the import name differ, which catches everyone
# once: you pip install python-docx, then you import docx.
try:
    from docx import Document
except ImportError:
    print("python-docx is not installed.")
    print("Install it with:  pip install python-docx")
    raise SystemExit(1)

if not DOCX_FILE.exists():
    print(f"{DOCX_FILE.name} is missing.")
    print("Run this first:  python create_sample_docx.py")
    raise SystemExit(1)

doc = Document(str(DOCX_FILE))


def paragraph_text(document: Document) -> str:
    """Join every paragraph into one string, the way the PDF demo does."""
    full_text = ""
    for para in document.paragraphs:
        full_text += para.text + "\n"
    return full_text


# --- 1. Paragraphs -------------------------------------------------------
# doc.paragraphs is body-level only. Text inside a table is NOT in here -
# tables are a separate collection, which is step 2.
full_text = paragraph_text(doc)
non_empty = [line for line in full_text.split("\n") if line.strip()]
print(f"--- 1. PARAGRAPHS ({len(non_empty)} of them) ---")
for line in non_empty:
    print(f"  {line}")

# --- 2. Tables - the thing PDFs cannot do -------------------------------
# A PDF would give you these cells as loose text with no way to tell which
# row or column they came from. Here the structure survives intact.
print(f"\n--- 2. TABLES ({len(doc.tables)}) ---")
for table_number, table in enumerate(doc.tables, start=1):
    print(f"Table {table_number}: {len(table.rows)} rows x {len(table.columns)} columns")
    for row in table.rows:
        cells = [cell.text for cell in row.cells]
        print(f"  {cells}")

# --- 3. Heading styles - the other thing PDFs lose ----------------------
# Every paragraph carries the style it was written with, so "which lines
# are headings?" is a question you can actually answer. In a PDF you would
# be guessing from capitalisation.
print("\n--- 3. HEADINGS ---")
for para in doc.paragraphs:
    if para.style.name == "Heading 1":
        print(f"HEADING: {para.text}")

# --- 4. Regex on the extracted text -------------------------------------
# Identical to the PDF demo. Once it is a string, the source format stops
# mattering - which is exactly what file 11's extract_text() exploits.
print("\n--- 4. REGEX ON THE TEXT ---")
print(f"Emails found: {re.findall(EMAIL_PATTERN, full_text)}")

# WHEN YOU GET A CHOICE, ASK FOR .DOCX
#   PDF  : flat string, tables lost, headings lost, scans give nothing.
#   DOCX : paragraphs, tables and styles all survive.
#   The catch is that resumes almost always arrive as PDFs, which is why
#   file 11 handles both and picks by file extension.
